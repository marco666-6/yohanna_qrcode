from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
OUT_PATH = OUT_DIR / "Panduan_Penjelasan_Project_Yohanna_Attendance_QRCode.docx"


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
PALE_YELLOW = "FFF6D5"
WHITE = "FFFFFF"
TEXT = "1F2937"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width=9360, indent=120):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def format_cell(cell, bold=False, fill=None, color=TEXT, size=9.5, align=None):
    if fill:
        set_cell_shading(cell, fill)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)
    for paragraph in cell.paragraphs:
        if align:
            paragraph.alignment = align
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.08
        for run in paragraph.runs:
            run.font.name = "Calibri"
            run.font.size = Pt(size)
            run.font.bold = bold
            run.font.color.rgb = RGBColor.from_string(color)


def add_table(doc, headers, rows, widths, font_size=9.3):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, text in enumerate(headers):
        cell = hdr.cells[idx]
        cell.text = text
        set_cell_width(cell, widths[idx])
        format_cell(cell, bold=True, fill=LIGHT_BLUE, color="0B2545", size=font_size)
    for row_data in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row_data):
            cells[idx].text = str(text)
            set_cell_width(cells[idx], widths[idx])
            format_cell(cells[idx], size=font_size)
    doc.add_paragraph()
    return table


def add_code_block(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table)
    cell = table.cell(0, 0)
    cell.text = "\n".join(lines)
    set_cell_shading(cell, "F8FAFC")
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    for p in cell.paragraphs:
        p.paragraph_format.space_after = Pt(0)
        for r in p.runs:
            r.font.name = "Consolas"
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor.from_string("111827")
    doc.add_paragraph()


def add_placeholder(doc, title, instruction):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, PALE_YELLOW)
    set_cell_margins(cell, top=180, bottom=180, start=180, end=180)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[TEMPATKAN SCREENSHOT: {title}]\n")
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string("7A5A00")
    run2 = p.add_run(instruction)
    run2.font.size = Pt(9.3)
    run2.font.color.rgb = RGBColor.from_string("5F4700")
    doc.add_paragraph()


def add_callout(doc, title, body, fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    p = cell.paragraphs[0]
    r = p.add_run(title + "\n")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string("0B2545")
    r.font.size = Pt(10.5)
    r2 = p.add_run(body)
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor.from_string(TEXT)
    doc.add_paragraph()


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_h1(doc, text):
    doc.add_heading(text, level=1)


def add_h2(doc, text):
    doc.add_heading(text, level=2)


def add_h3(doc, text):
    doc.add_heading(text, level=3)


def add_para(doc, text=""):
    p = doc.add_paragraph()
    p.add_run(text)
    return p


def setup_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for sec in doc.sections:
        sec.top_margin = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin = Inches(1)
        sec.right_margin = Inches(1)
        sec.header_distance = Inches(0.492)
        sec.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.167


def add_footer(doc):
    for section in doc.sections:
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = paragraph.add_run("Panduan Project Attendance QR Code - Yohanna")
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor.from_string("6B7280")


def build():
    OUT_DIR.mkdir(exist_ok=True)
    doc = Document()
    setup_styles(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Panduan Penjelasan Project Attendance QR Code")
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("0B2545")

    subtitle = doc.add_paragraph()
    subtitle.add_run("Dokumentasi implementasi, konfigurasi, alur kode, alur data, MVC, fitur utama, helper, dan proses otomatis untuk persiapan Sidang Akhir.").italic = True
    subtitle.paragraph_format.space_after = Pt(12)

    add_callout(
        doc,
        "Cara membaca dokumen ini",
        "Dokumen ini ditulis untuk pembaca yang belum terlalu nyaman dengan programming. Fokusnya bukan menghafal semua kode, tetapi memahami fungsi tiap bagian, alasan konfigurasi, dan alur kerja sistem dari pengguna sampai database.",
    )
    add_placeholder(
        doc,
        "Cover / halaman utama aplikasi",
        "Gunakan screenshot halaman login atau dashboard awal agar penguji langsung melihat konteks aplikasi.",
    )

    add_h1(doc, "1. Ringkasan Project")
    add_para(doc, "Project ini adalah aplikasi absensi karyawan berbasis Laravel 12. Sistem menggunakan QR Code harian untuk proses check-in dan check-out, memiliki tiga role utama yaitu Admin, HRD, dan Employee, serta menyediakan laporan, grafik statistik, notifikasi, ekspor Excel/PDF, dan pengajuan cuti.")
    add_table(
        doc,
        ["Aspek", "Penjelasan Singkat"],
        [
            ["Framework", "Laravel 12 sebagai backend MVC dan Blade sebagai tampilan server-side."],
            ["Frontend", "Vite, Tailwind CSS, Bootstrap icon/style pada view, serta JavaScript untuk scanner dan interaksi halaman."],
            ["Database", "MySQL dengan tabel users, shifts, qr_codes, attendances, leave_requests, notifications, activity_logs, sessions, dan cache."],
            ["Keamanan", "Login Laravel, hash password, middleware auth, middleware role, dan pengecekan akun aktif."],
            ["Output", "Data tampil di dashboard, riwayat, laporan, export Excel, export PDF, dan notifikasi/email."],
        ],
        [1800, 7560],
    )

    add_h2(doc, "1.1 Tujuan Sistem")
    add_bullets(doc, [
        "Mencatat kehadiran karyawan secara lebih cepat melalui QR Code.",
        "Membedakan hak akses antara Admin, HRD, dan Employee.",
        "Menyediakan laporan kehadiran yang dapat diekspor untuk kebutuhan administrasi.",
        "Memudahkan pengajuan dan persetujuan cuti atau ketidakhadiran.",
        "Membantu HRD memantau keterlambatan, cuti, ketidakhadiran, dan statistik absensi.",
    ])

    add_h2(doc, "1.2 Functional Requirement yang Dipenuhi")
    req_rows = [
        ["F001-F003", "Auth", "Login, logout, dan ubah password melalui AuthController dan route auth."],
        ["F004-F008", "Admin", "Kelola karyawan, QR Code, shift, toleransi terlambat, ekspor karyawan, serta force add/update absensi."],
        ["F009-F012", "Employee", "Scan QR, cek status hari ini, riwayat, notifikasi, dan pengajuan cuti."],
        ["F013-F016", "HRD", "Monitoring absensi, approve/reject cuti, ekspor laporan Excel/PDF, dan grafik statistik."],
    ]
    add_table(doc, ["Kode", "Area", "Implementasi"], req_rows, [1300, 1600, 6460])

    add_h1(doc, "2. Struktur Folder Penting")
    add_para(doc, "Laravel memakai struktur folder yang sudah tertata. Untuk sidang, tidak perlu menjelaskan semua file, tetapi perlu memahami file yang paling berpengaruh terhadap fitur.")
    add_table(
        doc,
        ["Folder/File", "Fungsi dalam Project"],
        [
            ["routes/web.php", "Peta alamat URL aplikasi. Di sini terlihat pembagian route Admin, HRD, Employee, attendance scan, dan QR Code API."],
            ["app/Http/Controllers", "Tempat controller. Controller menerima request, memvalidasi input, memanggil model/service, lalu mengirim hasil ke view."],
            ["app/Models", "Representasi tabel database seperti User, Attendance, Shift, QrCode, LeaveRequest, Notification, dan ActivityLog."],
            ["app/Services", "Tempat business logic yang dipisahkan dari controller, terutama QrCodeService dan EmployeeAttendanceService."],
            ["app/Helpers/helpers.php", "Kumpulan fungsi bantu global, misalnya format tanggal Indonesia, status badge, teks role, dan attendance window."],
            ["resources/views", "File Blade untuk tampilan login, dashboard, scanner, laporan, form cuti, dan halaman admin/HRD/employee."],
            ["database/migrations", "Definisi struktur tabel database."],
            ["database/seeders", "Data awal untuk demo, termasuk akun Admin, HRD, Employee, dan shift."],
            ["config/attendance.php", "Konfigurasi khusus absensi seperti masa aktif QR, window scan, toleransi terlambat, dan minimal jam kerja."],
            ["routes/console.php", "Scheduler otomatis untuk menjaga QR Code aktif sesuai jam absensi."],
        ],
        [2600, 6760],
    )
    add_placeholder(
        doc,
        "Struktur folder project di VS Code",
        "Screenshot Explorer VS Code yang menampilkan folder app, config, database, resources, routes, dan file composer/package.",
    )

    add_h1(doc, "3. Application Config")
    add_h2(doc, "3.1 Konfigurasi .env")
    add_para(doc, "File .env menyimpan pengaturan environment. File ini penting karena aplikasi yang sama bisa berjalan di laptop developer, server staging, atau server production dengan nilai konfigurasi berbeda.")
    add_code_block(doc, [
        'APP_NAME="PT Arung Laut Nusantara Attendance"',
        "APP_ENV=local",
        "APP_DEBUG=true",
        "APP_TIMEZONE=Asia/Jakarta",
        "APP_URL=http://localhost:8000",
        "",
        "DB_CONNECTION=mysql",
        "DB_HOST=127.0.0.1",
        "DB_PORT=3306",
        "DB_DATABASE=attendance_qrcode",
        "DB_USERNAME=root",
        "DB_PASSWORD=",
        "",
        "MAIL_MAILER=smtp",
        "MAIL_HOST=smtp.gmail.com",
        "MAIL_PORT=587",
        "MAIL_USERNAME=email@gmail.com",
        'MAIL_PASSWORD="app-pw"',
        "",
        "QR_CODE_EXPIRY_SECONDS=30",
        "QR_CODE_BEFORE_MINUTES=30",
        "QR_CODE_AFTER_MINUTES=45",
        "DEFAULT_LATE_TOLERANCE=15",
        "WORKING_HOURS_MINIMUM=8",
    ])
    add_table(
        doc,
        ["Konfigurasi", "Arti"],
        [
            ["APP_*", "Mengatur nama aplikasi, mode local/production, timezone, URL aplikasi, dan bahasa."],
            ["DB_*", "Mengatur koneksi ke database MySQL. Tanpa bagian ini, aplikasi tidak bisa membaca atau menyimpan data."],
            ["MAIL_*", "Dipakai untuk mengirim email notifikasi absensi dan status cuti."],
            ["QR_CODE_EXPIRY_SECONDS", "Durasi QR Code aktif. Pada project ini default-nya 30 detik."],
            ["QR_CODE_BEFORE_MINUTES", "QR mulai boleh dibuat sebelum jam target check-in/check-out. Default 30 menit."],
            ["QR_CODE_AFTER_MINUTES", "QR masih boleh tersedia setelah jam target. Default 45 menit."],
            ["DEFAULT_LATE_TOLERANCE", "Toleransi keterlambatan default dalam menit."],
            ["WORKING_HOURS_MINIMUM", "Minimal jam kerja yang menjadi acuan status kehadiran."],
        ],
        [2700, 6660],
    )

    add_h2(doc, "3.2 config/attendance.php")
    add_para(doc, "File config/attendance.php membaca nilai dari .env lalu menjadikannya konfigurasi Laravel yang bisa dipanggil melalui config('attendance.nama_key'). Ini membuat aturan absensi mudah diubah tanpa membongkar kode controller.")
    add_code_block(doc, [
        "'qr_code_expiry_seconds' => env('QR_CODE_EXPIRY_SECONDS', 30)",
        "'qr_code_before_minutes' => env('QR_CODE_BEFORE_MINUTES', 30)",
        "'qr_code_after_minutes' => env('QR_CODE_AFTER_MINUTES', 45)",
        "'default_late_tolerance' => env('DEFAULT_LATE_TOLERANCE', 15)",
        "'working_hours_minimum' => env('WORKING_HOURS_MINIMUM', 8)",
    ])

    add_h2(doc, "3.3 Middleware dan Bootstrap App")
    add_para(doc, "Pada bootstrap/app.php, project mendaftarkan middleware alias role, activity.log, dan user.active. Middleware adalah penjaga pintu sebelum request masuk ke controller.")
    add_table(
        doc,
        ["Middleware", "Peran"],
        [
            ["auth", "Memastikan user sudah login sebelum mengakses halaman tertentu."],
            ["role", "Memastikan hanya role tertentu yang bisa mengakses route, contoh admin hanya masuk /admin."],
            ["user.active", "Jika akun dinonaktifkan, user dipaksa logout dan diarahkan ke login."],
            ["activity.log", "Menyimpan aktivitas request user ke tabel activity_logs jika middleware dipakai."],
        ],
        [2200, 7160],
    )

    add_h1(doc, "4. Command untuk Setup dan Menjalankan Project")
    add_h2(doc, "4.1 Persiapan Dependency")
    add_code_block(doc, [
        "composer install",
        "npm install",
    ])
    add_para(doc, "composer install mengunduh dependency PHP/Laravel dari composer.json. npm install mengunduh dependency frontend seperti Vite, Tailwind CSS, Axios, dan Concurrently.")

    add_h2(doc, "4.2 Membuat dan Mengatur Environment")
    add_code_block(doc, [
        "Copy-Item .env.example .env",
        "php artisan key:generate",
    ])
    add_para(doc, "File .env dibuat dari .env.example, lalu application key dibuat agar fitur enkripsi Laravel seperti session dan token bekerja aman.")

    add_h2(doc, "4.3 Database dan Storage")
    add_code_block(doc, [
        "php artisan migrate:fresh --seed",
        "php artisan storage:link",
    ])
    add_para(doc, "migrate:fresh --seed membuat ulang tabel database dan mengisi data awal untuk demo. storage:link membuat folder public/storage agar file upload seperti lampiran cuti dapat diakses dari browser.")

    add_h2(doc, "4.4 Cache dan Menjalankan Aplikasi")
    add_code_block(doc, [
        "php artisan optimize:clear",
        "php artisan config:clear",
        "php artisan route:clear",
        "php artisan view:clear",
        "php artisan cache:clear",
        "",
        "composer run dev",
    ])
    add_para(doc, "composer run dev menjalankan beberapa proses sekaligus: server Laravel, queue listener, log viewer, scheduler, dan Vite. Ini praktis saat development karena backend, frontend, log, dan job otomatis berjalan bersama.")
    add_placeholder(
        doc,
        "Terminal composer run dev",
        "Screenshot terminal yang menunjukkan php artisan serve, queue, pail/log, schedule:work, dan npm run dev berjalan.",
    )

    add_h2(doc, "4.5 Command Otomatis QR Code")
    add_code_block(doc, [
        "php artisan attendance:generate-active-qr",
        "php artisan schedule:work",
    ])
    add_para(doc, "Command attendance:generate-active-qr memanggil QrCodeService untuk memastikan QR aktif tersedia saat window check-in atau check-out sedang terbuka. Scheduler menjalankannya setiap 10 detik dengan withoutOverlapping agar proses tidak bertabrakan.")

    add_h1(doc, "5. Konsep MVC pada Project")
    add_para(doc, "MVC adalah pola pemisahan tanggung jawab: Model mengurus data, View mengurus tampilan, dan Controller mengatur alur request. Analogi sederhananya: Model adalah arsip data, View adalah layar yang dilihat user, Controller adalah petugas yang menerima permintaan lalu mengambil data yang tepat.")
    add_table(
        doc,
        ["Bagian MVC", "Contoh di Project", "Penjelasan"],
        [
            ["Model", "User, Attendance, Shift, QrCode", "Mewakili tabel database dan relationship antar data."],
            ["View", "resources/views/employee/scanner.blade.php", "Menampilkan UI scanner QR dan informasi status absensi."],
            ["Controller", "AttendanceController@scan", "Menerima hasil scan, validasi QR, membuat/update attendance, lalu mengirim response."],
            ["Service", "QrCodeService, EmployeeAttendanceService", "Bukan bagian MVC klasik, tetapi membantu memisahkan business logic agar controller tidak terlalu panjang."],
        ],
        [1800, 2900, 4660],
        font_size=8.8,
    )

    add_h2(doc, "5.1 Contoh Alur MVC: Employee Scan QR")
    add_numbers(doc, [
        "Employee membuka halaman scanner pada route /employee/scanner.",
        "View scanner menampilkan kamera/QR scanner dan status absensi hari ini.",
        "Saat QR berhasil dibaca, browser mengirim POST ke /attendance/scan.",
        "AttendanceController memvalidasi kode QR, shift, window waktu, dan status absensi user.",
        "Model Attendance menyimpan check-in atau check-out ke database.",
        "Notification dan email dibuat untuk memberi tahu hasil absensi.",
        "Response JSON dikirim kembali ke browser agar user melihat status berhasil atau gagal.",
    ])
    add_placeholder(
        doc,
        "Halaman scanner employee",
        "Screenshot halaman scanner saat QR aktif muncul atau saat kamera siap scan.",
    )

    add_h1(doc, "6. Code-flow dan Data-flow")
    add_h2(doc, "6.1 Login dan Redirect Role")
    add_para(doc, "Login ditangani AuthController. User memasukkan email dan password, Laravel memeriksa kredensial, lalu sistem melihat role user untuk menentukan dashboard tujuan.")
    add_code_block(doc, [
        "User -> route /login -> AuthController@login",
        "Validasi email/password -> Auth::attempt",
        "Jika berhasil -> regenerate session",
        "role admin -> /admin/dashboard",
        "role hrd -> /hrd/dashboard",
        "role employee -> /employee/dashboard",
    ])
    add_placeholder(
        doc,
        "Halaman login dan contoh redirect dashboard",
        "Bisa pakai dua screenshot kecil: form login dan dashboard sesuai role setelah login.",
    )

    add_h2(doc, "6.2 Generate QR Code")
    add_para(doc, "QR Code tidak dibuat sembarangan. Sistem melihat shift aktif, menentukan apakah waktu check-in/check-out sedang terbuka, lalu membuat QR aktif jika belum ada QR yang valid.")
    add_code_block(doc, [
        "Scheduler setiap 10 detik",
        "-> attendance:generate-active-qr",
        "-> QrCodeService@ensureCurrentWindowQRCodes",
        "-> loop semua shift aktif",
        "-> getAttendanceWindow(shift, check_in/check_out)",
        "-> jika window terbuka: ensureActiveCode",
        "-> jika window tertutup: nonaktifkan QR lama",
    ])
    add_placeholder(
        doc,
        "Halaman Admin QR Code",
        "Screenshot halaman QR Code admin yang menampilkan QR check-in/check-out, shift, masa berlaku, atau tombol generate.",
    )

    add_h2(doc, "6.3 Scan QR dan Penyimpanan Attendance")
    add_para(doc, "Saat employee scan QR, data yang mengalir bukan gambar QR-nya, melainkan string code unik dari tabel qr_codes. Controller mencocokkan code tersebut dengan QR aktif di database.")
    add_table(
        doc,
        ["Tahap", "Data yang Diproses", "Hasil"],
        [
            ["Scan", "code QR dari browser", "Dikirim ke AttendanceController@scan."],
            ["Validasi", "code, type, shift_id, expires_at, is_active", "Menentukan QR valid atau tidak."],
            ["Check-in", "user_id, date, check_in, check_in_qr_id, shift_id", "Membuat row attendance baru atau update row hari ini."],
            ["Check-out", "check_out, check_out_qr_id, total_hours", "Melengkapi absensi yang sudah check-in."],
            ["Status", "jam check-in, toleransi shift, total jam", "Menghasilkan on_time, late, incomplete, atau absent."],
            ["Notifikasi", "user, type, status, time", "Muncul di halaman notification dan dapat dikirim via email."],
        ],
        [1500, 3860, 4000],
        font_size=8.8,
    )

    add_h2(doc, "6.4 Pengajuan Cuti")
    add_para(doc, "Employee membuat pengajuan cuti dari form. Data masuk ke tabel leave_requests dengan status pending. HRD/Admin kemudian membuka daftar pengajuan, melihat detail, lalu approve atau reject. Saat status berubah, sistem mengisi reviewed_by, reviewed_at, dan review_notes.")
    add_placeholder(
        doc,
        "Form pengajuan cuti dan halaman approval HRD",
        "Masukkan screenshot form employee serta detail approval/rejection di HRD agar alur pending -> approved/rejected terlihat.",
    )

    add_h2(doc, "6.5 Laporan dan Export")
    add_para(doc, "HRD dapat memfilter laporan kehadiran berdasarkan tanggal, lalu mengekspor ke Excel atau PDF. Export Excel memakai Maatwebsite Excel, sedangkan PDF memakai DomPDF.")
    add_bullets(doc, [
        "AttendanceExport mengambil data attendance beserta user dan shift.",
        "EmployeesExport mengambil data user employee beserta shift.",
        "View hrd/exports/attendance-pdf.blade.php menjadi template PDF laporan.",
        "Helper formatDate, formatTime, dan getStatusText dipakai agar laporan mudah dibaca dalam bahasa Indonesia.",
    ])
    add_placeholder(
        doc,
        "Laporan HRD dan hasil export",
        "Screenshot halaman attendance report serta contoh file Excel/PDF yang berhasil dibuat.",
    )

    add_h1(doc, "7. Database dan Relasi Data")
    add_para(doc, "Database adalah tempat sistem menyimpan kondisi nyata: siapa user-nya, shift apa yang dimiliki, QR apa yang aktif, absensi apa yang tercatat, dan cuti apa yang diajukan.")
    add_table(
        doc,
        ["Tabel", "Isi Utama", "Relasi Penting"],
        [
            ["users", "Akun, role, data karyawan, shift_id, status aktif", "belongsTo Shift; hasMany Attendance, LeaveRequest, Notification."],
            ["shifts", "Nama shift, jam mulai, jam selesai, toleransi terlambat", "hasMany User, Attendance, QrCode."],
            ["qr_codes", "Kode unik, type check_in/check_out, shift, expired, aktif/tidak", "belongsTo Shift; dipakai Attendance sebagai check_in_qr/check_out_qr."],
            ["attendances", "Tanggal, check-in, check-out, status, total jam, catatan", "belongsTo User, Shift, QrCode."],
            ["leave_requests", "Jenis cuti, tanggal, alasan, status, reviewer", "belongsTo User dan reviewer dari tabel users."],
            ["notifications", "Judul, pesan, tipe, status baca", "belongsTo User."],
            ["activity_logs", "Aksi request user, IP, user agent", "belongsTo User."],
        ],
        [1800, 3600, 3960],
        font_size=8.4,
    )
    add_placeholder(
        doc,
        "ERD database",
        "Buat diagram relasi tabel users, shifts, qr_codes, attendances, leave_requests, notifications, activity_logs. Bisa dari draw.io/dbdiagram.",
    )

    add_h1(doc, "8. Fitur Utama per Role")
    add_h2(doc, "8.1 Admin")
    add_bullets(doc, [
        "Dashboard ringkasan jumlah karyawan, absensi, cuti, dan data operasional.",
        "CRUD karyawan: tambah, ubah, hapus/nonaktifkan, detail, dan export Excel.",
        "CRUD shift kerja dan toleransi keterlambatan.",
        "Manajemen attendance termasuk force add dan update data absensi.",
        "Generate dan monitoring QR Code harian.",
        "Melihat dan memproses leave request jika diberi route admin.",
    ])
    add_placeholder(doc, "Dashboard Admin", "Screenshot dashboard admin dengan kartu statistik dan navigasi utama.")
    add_placeholder(doc, "CRUD Karyawan", "Screenshot list karyawan, form tambah/edit, dan detail karyawan.")
    add_placeholder(doc, "Manajemen Shift", "Screenshot halaman shift yang menampilkan shift pagi/siang/malam dan toleransi terlambat.")

    add_h2(doc, "8.2 Employee")
    add_bullets(doc, [
        "Dashboard pribadi untuk melihat status absensi hari ini.",
        "Scanner QR untuk check-in dan check-out.",
        "Riwayat absensi pribadi dengan filter.",
        "Notifikasi berhasil/gagal absensi.",
        "Form pengajuan cuti atau ketidakhadiran.",
        "Profile dan ubah password pribadi.",
    ])
    add_placeholder(doc, "Dashboard Employee", "Screenshot dashboard employee yang menunjukkan status hari ini, next action, dan notice.")
    add_placeholder(doc, "Riwayat Absensi Employee", "Screenshot attendance history employee dengan tanggal, status, check-in, dan check-out.")

    add_h2(doc, "8.3 HRD")
    add_bullets(doc, [
        "Dashboard monitoring kehadiran.",
        "Laporan attendance dengan filter tanggal dan status.",
        "Approve/reject leave request dengan catatan review.",
        "Export laporan ke Excel dan PDF.",
        "Statistik/grafik attendance untuk melihat tren kehadiran, terlambat, cuti, dan tidak hadir.",
        "Menambahkan notes pada data attendance tertentu.",
    ])
    add_placeholder(doc, "Dashboard HRD dan Statistik", "Screenshot grafik statistik HRD dan ringkasan monitoring.")
    add_placeholder(doc, "Approval Cuti HRD", "Screenshot detail leave request dengan tombol approve/reject dan catatan.")

    add_h1(doc, "9. Helper dan Automatic Stuff")
    add_h2(doc, "9.1 Helper")
    add_para(doc, "Helper adalah fungsi bantu global agar kode tidak berulang di banyak controller atau view. Contohnya, daripada menulis format tanggal Indonesia di setiap halaman, project cukup memanggil formatDate().")
    add_table(
        doc,
        ["Helper", "Kegunaan"],
        [
            ["formatDate / formatTime / formatDateTime", "Mengubah tanggal dan jam menjadi format yang mudah dibaca user Indonesia."],
            ["getStatusBadge / getStatusText", "Mengubah status attendance menjadi teks Indonesia dan badge warna."],
            ["getLeaveTypeText / getLeaveStatusText", "Mengubah kode leave request menjadi label manusiawi."],
            ["getRoleText", "Mengubah role admin/hrd/employee menjadi label tampilan."],
            ["calculateWorkingDays", "Menghitung hari kerja antara dua tanggal tanpa weekend."],
            ["getAttendanceWindow", "Menghitung kapan window check-in/check-out dibuka dan ditutup."],
            ["attendancePercentage", "Menghitung persentase statistik dengan aman jika total data nol."],
            ["attendanceStatusIcon", "Memilih icon status attendance untuk UI."],
        ],
        [3000, 6360],
        font_size=8.9,
    )

    add_h2(doc, "9.2 Scheduler Otomatis")
    add_para(doc, "Scheduler adalah fitur Laravel untuk menjalankan command otomatis berdasarkan jadwal. Pada project ini, scheduler menjalankan command QR Code setiap 10 detik.")
    add_code_block(doc, [
        "Schedule::command('attendance:generate-active-qr --quiet-output')",
        "    ->everyTenSeconds()",
        "    ->withoutOverlapping();",
    ])
    add_callout(
        doc,
        "Poin penting untuk sidang",
        "Jika ditanya kenapa QR perlu otomatis: karena QR hanya valid pada window tertentu dan expired cepat. Scheduler memastikan user tidak perlu menunggu admin menekan tombol manual setiap saat.",
        fill="EAF4FF",
    )

    add_h2(doc, "9.3 Service Layer")
    add_table(
        doc,
        ["Service", "Tanggung Jawab"],
        [
            ["QrCodeService", "Membuat kode unik, menonaktifkan QR lama, menjaga QR aktif selama window absensi terbuka, dan regenerate QR untuk shift tertentu."],
            ["EmployeeAttendanceService", "Membangun konteks dashboard/scanner employee: shift, attendance hari ini, next action, window, QR aktif, notice, dan status tombol absensi."],
        ],
        [2800, 6560],
    )

    add_h1(doc, "10. Keamanan dan Validasi")
    add_bullets(doc, [
        "Password disimpan menggunakan Hash::make, bukan plain text.",
        "Route diproteksi auth agar halaman utama hanya bisa dibuka setelah login.",
        "Route per role diproteksi RoleMiddleware agar user tidak bisa masuk ke dashboard role lain.",
        "CheckUserActive memaksa logout akun yang sudah dinonaktifkan.",
        "Form penting memakai validasi request sebelum data masuk database.",
        "QR Code dicek is_active, expires_at, shift, dan type sebelum diterima sebagai absensi.",
        "Attendance memakai unique user_id + date agar satu user tidak punya dua data absensi untuk tanggal yang sama.",
    ])

    add_h1(doc, "11. Data Seeder untuk Demo Sidang")
    add_para(doc, "Seeder membantu menyiapkan data awal sehingga aplikasi langsung bisa didemokan setelah migration.")
    add_table(
        doc,
        ["Akun", "Email", "Password", "Role"],
        [
            ["Administrator", "admin@arunglaut.com", "admin123", "admin"],
            ["HR Manager", "hrd@arunglaut.com", "hrd123", "hrd"],
            ["Budi Santoso", "budi@arunglaut.com", "employee123", "employee"],
            ["Siti Rahayu", "siti@arunglaut.com", "employee123", "employee"],
            ["Ahmad Hidayat", "ahmad@arunglaut.com", "employee123", "employee"],
        ],
        [2200, 3100, 1900, 2160],
        font_size=8.8,
    )
    add_callout(
        doc,
        "Catatan keamanan",
        "Akun dan password seed boleh dipakai untuk demo local. Untuk production, password default harus diganti dan APP_DEBUG harus false.",
        fill="FFF6D5",
    )

    add_h1(doc, "12. Checklist Screenshot untuk Melengkapi Dokumen")
    add_table(
        doc,
        ["No", "Screenshot yang Disarankan", "Tujuan Penjelasan"],
        [
            ["1", "Halaman login", "Menunjukkan titik masuk semua role."],
            ["2", "Dashboard Admin", "Menunjukkan ringkasan kontrol admin."],
            ["3", "CRUD karyawan", "Membuktikan F004 dan F007."],
            ["4", "Manajemen shift", "Membuktikan F006."],
            ["5", "Halaman QR Code admin", "Membuktikan F005 dan konsep QR harian."],
            ["6", "Scanner employee", "Membuktikan F009."],
            ["7", "Status dan riwayat absensi employee", "Membuktikan F010-F011."],
            ["8", "Pengajuan cuti employee", "Membuktikan F012."],
            ["9", "Approval cuti HRD", "Membuktikan F013."],
            ["10", "Laporan attendance HRD", "Membuktikan F014-F015."],
            ["11", "Grafik statistik HRD", "Membuktikan F016."],
            ["12", "Terminal composer run dev / scheduler", "Menjelaskan automatic QR generation."],
            ["13", "ERD database", "Menjelaskan relasi data dan data-flow."],
        ],
        [700, 3300, 5360],
        font_size=8.5,
    )

    add_h1(doc, "13. Contoh Jawaban Singkat untuk Sidang")
    add_h3(doc, "Kenapa memakai QR Code?")
    add_para(doc, "Karena QR Code mempercepat proses absensi dan mengurangi input manual. Pada sistem ini QR juga dibuat aktif hanya pada window waktu tertentu dan expired dalam durasi singkat, sehingga lebih aman dibanding kode statis.")
    add_h3(doc, "Apa perbedaan Admin, HRD, dan Employee?")
    add_para(doc, "Admin fokus pada konfigurasi dan data master seperti karyawan, shift, QR Code, dan koreksi absensi. HRD fokus pada monitoring, approval cuti, laporan, dan statistik. Employee fokus pada absensi pribadi, riwayat, notifikasi, dan pengajuan cuti.")
    add_h3(doc, "Bagaimana sistem menentukan terlambat?")
    add_para(doc, "Sistem membandingkan jam check-in dengan jam mulai shift dan nilai late_tolerance. Jika check-in melewati batas toleransi, status menjadi late; jika masih dalam toleransi, status menjadi on_time.")
    add_h3(doc, "Apa fungsi scheduler?")
    add_para(doc, "Scheduler menjalankan command otomatis untuk memastikan QR Code aktif tersedia saat window check-in atau check-out terbuka. Jadi proses QR tidak bergantung sepenuhnya pada klik manual admin.")
    add_h3(doc, "Kenapa ada service layer?")
    add_para(doc, "Service layer membuat logic penting lebih rapi dan reusable. Contohnya logic QR Code ditempatkan di QrCodeService, sehingga controller tidak terlalu panjang dan lebih mudah diuji atau dipahami.")

    add_h1(doc, "14. Catatan Akhir untuk Yohanna")
    add_para(doc, "Saat mempresentasikan project, fokuskan cerita pada masalah, solusi, dan alur data. Tidak perlu menghafal seluruh kode. Yang penting adalah memahami: user login, role menentukan hak akses, QR dibuat sesuai shift dan waktu, hasil scan masuk ke attendance, HRD/Admin dapat memonitor dan mengekspor laporan, serta helper/service/scheduler membuat sistem lebih rapi dan otomatis.")
    add_callout(
        doc,
        "Strategi presentasi",
        "Mulai dari demo fitur, lalu jelaskan alur teknisnya. Jika penguji bertanya kode, arahkan ke file yang relevan: routes/web.php untuk URL, controller untuk proses, model/migration untuk database, service/helper untuk logic pendukung.",
        fill="EAF7EA",
    )

    add_footer(doc)
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build()
